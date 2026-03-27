from docx import Document
import os

def extract_word_content(docx_path):
    """提取Word文档内容"""
    try:
        doc = Document(docx_path)
        
        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # 提取所有表格内容
        tables_data = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_data.append(table_data)
        
        return paragraphs, tables_data
    
    except Exception as e:
        print(f"读取文档出错: {e}")
        return None, None

def save_content_to_file(paragraphs, tables_data, output_file):
    """将提取的内容保存到文件"""
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("=" * 50 + "\n")
        f.write("文档段落内容\n")
        f.write("=" * 50 + "\n\n")
        
        for i, para in enumerate(paragraphs, 1):
            f.write(f"[{i}] {para}\n\n")
        
        f.write("\n" + "=" * 50 + "\n")
        f.write("文档表格内容\n")
        f.write("=" * 50 + "\n")
        
        for table_idx, table in enumerate(tables_data, 1):
            f.write(f"\n--- 表格 {table_idx} ---\n")
            for row in table:
                f.write(" | ".join(row) + "\n")

def main():
    docx_file = "铸信+样本+2026.3.docx"
    output_file = "extracted_content.txt"
    
    if os.path.exists(docx_file):
        print(f"正在读取文档: {docx_file}")
        paragraphs, tables_data = extract_word_content(docx_file)
        
        if paragraphs or tables_data:
            save_content_to_file(paragraphs, tables_data, output_file)
            print(f"内容已保存到: {output_file}")
            
            # 打印部分内容预览
            print("\n--- 内容预览 ---")
            if paragraphs:
                print("\n前5段内容:")
                for i, para in enumerate(paragraphs[:5], 1):
                    print(f"[{i}] {para[:100]}..." if len(para) > 100 else f"[{i}] {para}")
            
            if tables_data:
                print(f"\n发现 {len(tables_data)} 个表格")
                for i, table in enumerate(tables_data[:2], 1):  # 只显示前2个表格
                    print(f"\n表格 {i} (共{len(table)}行):")
                    for row in table[:3]:  # 每个表格只显示前3行
                        print(" | ".join(row)[:80] + "..." if len(" | ".join(row)) > 80 else " | ".join(row))
                    if len(table) > 3:
                        print("...")
        else:
            print("未能提取到任何内容")
    else:
        print(f"文件不存在: {docx_file}")

if __name__ == "__main__":
    main()
