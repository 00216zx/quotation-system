from docx import Document

try:
    doc = Document('铸信+样本+2026.3.docx')
    full_text = []
    
    print("=== 文档内容 ===\n")
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            print(para.text)
    
    print("\n=== 表格内容 ===\n")
    
    for i, table in enumerate(doc.tables):
        print(f"\n--- 表格 {i+1} ---\n")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(' | '.join(row_text))
            
except Exception as e:
    print(f"错误: {e}")
